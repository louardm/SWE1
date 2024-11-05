from deepdiff import DeepDiff
from docx import Document
import json
import os
from PyQt5.QtWidgets import QMessageBox


def save_json_diff_to_doc(file_a_path, file_b_path, output_path="json_differences.doc"):
    try:
        # Load JSON content
        with open(file_a_path, 'r') as file_a, open(file_b_path, 'r') as file_b:
            json_a = json.load(file_a)
            json_b = json.load(file_b)

        # Compare the JSON files
        diff = DeepDiff(json_a, json_b, ignore_order=True)

        # Check if there are any differences
        if not diff:
            return None  # No differences

        # Create a .doc file with the differences
        doc = Document()
        doc.add_heading('JSON Differences', level=1)

        for key, value in diff.items():
            doc.add_paragraph(f"{key}:")
            doc.add_paragraph(json.dumps(value, indent=4), style="Quote")

        # Save the document
        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error saving JSON diff to .doc: {e}")
        return None
